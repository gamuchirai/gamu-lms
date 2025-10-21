from docx import Document
from docx.shared import Pt

OUTPUT_PATH = r"c:\Users\DoddleLearnZW\Documents\code\Gamuchirai Kundhlande Dzidza LMS\public\How_Web_Form_Security_Prevents_Intrusion.docx"

content = {
    "title": "How Web Form Security Prevents Intrusion",
    "sections": [
        {
            "heading": "Overview",
            "body": (
                "Web form security reduces attack surface and stops automated and manual intrusions by validating input, "
                "authenticating users, protecting sessions, and sanitizing output."
            ),
        },
        {
            "heading": "Input Validation & Sanitization",
            "body": (
                "Server-side validation (required) and client-side checks (UX) prevent malformed or dangerous data. "
                "Sanitization removes/encodes characters that could trigger SQL injection or XSS. In this codebase, "
                "use prepared statements and parameterized queries to avoid SQL injection, and escape output when rendering."
            ),
        },
        {
            "heading": "Authentication & Email Verification",
            "body": (
                "Require verified accounts and enforce strong password handling. Email verification prevents automated account abuse. "
                "Implement account lockouts and rate-limiting on login and resend verification endpoints to stop brute force and enumeration."
            ),
        },
        {
            "heading": "CSRF Protection",
            "body": (
                "Use anti-CSRF tokens for state-changing forms (login, register, update). Tokens ensure requests originate from legitimate pages, "
                "preventing cross-site request forgery attacks."
            ),
        },
        {
            "heading": "Session Management & Secure Cookies",
            "body": (
                "Use server-side session checks, regenerate session IDs after privilege changes, and set cookies with HttpOnly and Secure flags. "
                "Validate session on sensitive operations and log out on suspicious activity."
            ),
        },
        {
            "heading": "Rate Limiting & Logging",
            "body": (
                "Apply rate limits to critical endpoints (login, resend verification) and log suspicious events for analysis. "
                "Logs help detect repeated intrusion attempts and support incident response."
            ),
        },
        {
            "heading": "Summary",
            "body": (
                "Combining validation, sanitization, authentication, CSRF protection, secure session handling, and monitoring "
                "creates layered defenses that prevent most common intrusion vectors in web forms."
            ),
        },
    ],
}


def make_doc(path):
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    doc.add_heading(content['title'], level=1)

    for sec in content['sections']:
        doc.add_heading(sec['heading'], level=2)
        p = doc.add_paragraph(sec['body'])
        p.style = doc.styles['Normal']

    doc.save(path)


if __name__ == '__main__':
    make_doc(OUTPUT_PATH)
    print(f"Wrote: {OUTPUT_PATH}")
